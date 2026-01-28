"""
Document loader interface and lightweight implementation.
Uses pypdf, python-docx, and native Python for basic formats.
"""

from abc import ABC, abstractmethod
from typing import List, Union
from pathlib import Path
import pypdf
import docx
from domain.entities import DocumentChunk
from core.exceptions import DocumentProcessingError
from core.logging import get_logger

logger = get_logger(__name__)


class IDocumentLoader(ABC):
    """Interface for document loaders."""
    
    @abstractmethod
    def load(self, file_path: Union[str, Path]) -> List[DocumentChunk]:
        """
        Load and parse a document into chunks.
        
        Args:
            file_path: Path to the document
            
        Returns:
            List of DocumentChunk objects
        """
        pass


class LightweightDocumentLoader(IDocumentLoader):
    """Document loader using lightweight, pure-Python libraries."""
    
    SUPPORTED_EXTENSIONS = {'.txt', '.md', '.pdf', '.docx', '.doc'}
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB
    
    def __init__(self):
        """Initialize the document loader."""
        logger.info("Initialized Lightweight document loader")
    
    def load(self, file_path: Union[str, Path]) -> List[DocumentChunk]:
        """Load and parse a document into chunks."""
        try:
            file_path = Path(file_path)
            
            # Validate file exists
            if not file_path.exists():
                raise DocumentProcessingError(f"File not found: {file_path}")

            # Guardrail: Check file size
            file_size = file_path.stat().st_size
            if file_size > self.MAX_FILE_SIZE:
                raise DocumentProcessingError(
                    f"File too large: {file_size / (1024*1024):.1f}MB. "
                    f"Max allowed size: {self.MAX_FILE_SIZE / (1024*1024)}MB"
                )
            
            # Validate file type
            suffix = file_path.suffix.lower()
            if suffix not in self.SUPPORTED_EXTENSIONS:
                raise DocumentProcessingError(
                    f"Unsupported file type: {suffix}. "
                    f"Supported types: {', '.join(self.SUPPORTED_EXTENSIONS)}"
                )
            
            logger.info(f"Loading document ({suffix}): {file_path}")
            
            # Dispatch to specific loader based on extension
            if suffix in ('.txt', '.md'):
                return self._load_text(file_path)
            elif suffix == '.pdf':
                return self._load_pdf(file_path)
            elif suffix in ('.docx', '.doc'):
                return self._load_docx(file_path)
            else:
                raise DocumentProcessingError(f"Logic error: extension {suffix} not handled")
                
        except Exception as e:
            if isinstance(e, DocumentProcessingError):
                raise
            error_msg = f"Failed to load document {file_path}: {str(e)}"
            logger.error(error_msg)
            raise DocumentProcessingError(error_msg)

    def _load_text(self, file_path: Path) -> List[DocumentChunk]:
        """Load character-based text files."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        # Split by paragraphs to provide better fragments to the chunker
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        
        return [
            DocumentChunk(
                text=para,
                chunk_index=i,
                metadata={
                    'source': str(file_path),
                    'file_type': file_path.suffix.lower(),
                    'method': 'native_text'
                }
            )
            for i, para in enumerate(paragraphs)
        ]

    def _load_pdf(self, file_path: Path) -> List[DocumentChunk]:
        """Load PDF using pypdf."""
        chunks = []
        try:
            reader = pypdf.PdfReader(file_path)
            full_text = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    full_text.append(page_text.strip())
            
            if not full_text:
                raise DocumentProcessingError("PDF contains no readable text.")

            # Join all pages with spacing
            combined_text = "\n\n".join(full_text)
            
            # Split by paragraphs to provide better fragments to the chunker
            paragraphs = [p.strip() for p in combined_text.split('\n\n') if p.strip()]
            
            return [
                DocumentChunk(
                    text=para,
                    chunk_index=i,
                    metadata={
                        'source': str(file_path),
                        'file_type': '.pdf',
                        'pages': len(reader.pages),
                        'method': 'pypdf'
                    }
                )
                for i, para in enumerate(paragraphs)
            ]
        except Exception as e:
            raise DocumentProcessingError(f"Error parsing PDF: {str(e)}")

    def _load_docx(self, file_path: Path) -> List[DocumentChunk]:
        """Load Word documents using python-docx."""
        try:
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text and para.text.strip():
                    full_text.append(para.text.strip())
            
            if not full_text:
                # Check tables if no paragraph text found
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text and cell.text.strip():
                                full_text.append(cell.text.strip())
            
            if not full_text:
                raise DocumentProcessingError("Word document contains no readable text.")

            combined_text = "\n\n".join(full_text)
            
            # Split by paragraphs to provide better fragments to the chunker
            paragraphs = [p.strip() for p in combined_text.split('\n\n') if p.strip()]
            
            return [
                DocumentChunk(
                    text=para,
                    chunk_index=i,
                    metadata={
                        'source': str(file_path),
                        'file_type': file_path.suffix.lower(),
                        'method': 'python-docx'
                    }
                )
                for i, para in enumerate(paragraphs)
            ]
        except Exception as e:
            raise DocumentProcessingError(f"Error parsing Word document: {str(e)}")
